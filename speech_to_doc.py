"""
    This script takes in speech from the user and transcripts it to a docx file inside
    local folder unless otherwise specified.  
"""
# speech_to_document.py
import speech_recognition as sr
from docx import Document
import os


def transcribe_speech_to_docx(doc_path=None):
    # Initialize the recognizer
    r = sr.Recognizer()

    # If doc_path is None or empty, ask for input or set a default
    if not doc_path:
        # Example: Prompt for path
        doc_path = input("Enter the document path including filename (or just filename): ")


    # Check if the document exists, if not create a new one
    if not os.path.exists(doc_path):
        doc = Document()
        doc.save(doc_path)

    # Use the default microphone as the audio source
    with sr.Microphone() as source:
        print("Please speak something... (Press Ctrl+C to stop)")

        # Adjust the recognizer sensitivity to ambient noise
        r.adjust_for_ambient_noise(source, duration=1)

        try:
            while True:
                print("Listening...")
                audio = r.listen(source)
                
                # Recognize the speech
                try:
                    text = r.recognize_google(audio)
                    print("You said: " + text)

                    # Append the text to the Word document
                    doc = Document(doc_path)
                    doc.add_paragraph(text)
                    doc.save(doc_path)

                except sr.UnknownValueError:
                    print("Could not understand the audio")
                except sr.RequestError as e:
                    print(f"Could not request results; {e}")
        except KeyboardInterrupt:
            # Stop the program when Ctrl+C is pressed
            print("Stopping speech recognition...")

