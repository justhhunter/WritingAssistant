"""
    This script edits a file using the gpt api using
    a special prompt for editing lab reports
"""


import openai
from docx import Document
import os


def edit_document(doc_path, openai_api_key):
    
    

    # get text from document
    doc = Document(doc_path)
    full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip() != ''])

    # Define the GPT prompt for editing
    prompt_text = f"Please edit the following text for clarity, grammar, and professionalism, while maintaining the original meaning and style:\n\n{full_text}"

    # Send the prompt to GPT
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.3,
        max_tokens=4096,  
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt_text}
        ]
    )

    edited_text = response.choices[0].message.content

    # Save the edited text to a new document
    new_doc_path = os.path.splitext(doc_path)[0] + "_edited.docx"
    new_doc = Document()
    new_doc.add_paragraph(edited_text)
    new_doc.save(new_doc_path)

    return new_doc_path

